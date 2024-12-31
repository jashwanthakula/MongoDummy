import streamlit as st
import pymongo
import nltk
from nltk.corpus import wordnet as wn
from nltk.tokenize import word_tokenize
from docx import Document
from fpdf import FPDF
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from io import BytesIO

# Ensure NLTK resources are downloaded
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/wordnet')
except LookupError:
    nltk.download('wordnet')

nltk.download('punkt', quiet=True)
# Connect to MongoDB Atlas
def connect_to_mongodb():
    client = pymongo.MongoClient("mongodb+srv://jashwanthakula26:majormongo0745@cluster0.uqvju.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0")
    db = client["yoga_asana"]  # Replace with your database name
    return db


# Function to extract keywords and find synonyms
def extract_keywords(health_issue):
    # Tokenize the input sentence
    words = nltk.word_tokenize(health_issue)
    keywords = []
    
    # Extract keywords (you can add more criteria for filtering)
    for word in words:
        if word.isalpha():  # Filter out punctuation
            keywords.append(word.lower())
    
    synonyms = []
    
    # Find synonyms for each keyword
    for keyword in keywords:
        for syn in wn.synsets(keyword):
            for lemma in syn.lemmas():
                synonyms.append(lemma.name())  # Get the synonym name
    return list(set(synonyms))  # Return unique synonyms


# Function to fetch matching asanas from the database
def get_matching_asanas(age, gender, health_issue):
    db = connect_to_mongodb()
    collection = db["yoga_asanas"]  # Replace with your collection name

    # Extract keywords and synonyms from health issue
    keywords = extract_keywords(health_issue)

    # MongoDB query
    query = {
        "$and": [
            {"$or": [{"gender": gender}, {"gender": "All"}]},
            {"min_age": {"$lte": age}},
            {"$or": [{"health_issue": {"$regex": keyword, "$options": "i"}} for keyword in keywords]}
        ]
    }

    # Fetch data from MongoDB
    results = collection.find(query)
    
    # Return results as a list of tuples
    return [(result["asana_name"], result["health_issue"]) for result in results]


# Function to create a Word document in memory
def create_word_document(recommendations):
    doc = Document()
    doc.add_heading("Yoga Recommendations", level=1)
    doc.add_paragraph("Here are your personalized yoga recommendations:")
    for rec in recommendations:
        doc.add_paragraph(f"- {rec}")
    # Save to BytesIO instead of a file
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)  # Reset stream position for reading
    return file_stream


# Function to create a PDF document in memory
def create_pdf_document(recommendations):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Yoga Recommendations", ln=True, align='C')
    pdf.ln(10)  # Add a line break
    pdf.set_font("Arial", size=10)
    pdf.cell(200, 10, txt="Here are your personalized yoga recommendations:", ln=True)
    pdf.ln(10)
    for rec in recommendations:
        pdf.cell(0, 10, txt=f"- {rec}", ln=True)
    
    # Generate PDF content as a string
    pdf_content = pdf.output(dest='S').encode('latin1')

    # Write the content to a BytesIO stream
    file_stream = BytesIO(pdf_content)
    file_stream.seek(0)  # Reset stream position for reading
    return file_stream


# Function to send recommendations via email with attachment
def send_email(user_email, file_stream, file_name):
    try:
        # Email configuration
        sender_email = "jahnaviproject7@gmail.com"
        sender_password = "iwlw hosy ifat lfmd"  # Use an App Password if using Gmail
        subject = "Your Personalized Yoga Recommendations"

        # Create email content
        message = MIMEMultipart()
        message['From'] = sender_email
        message['To'] = user_email
        message['Subject'] = subject

        body = "Hello,\n\nPlease find attached your personalized yoga recommendations.\n\nStay Healthy!"
        message.attach(MIMEText(body, 'plain'))

        # Attach the file
        part = MIMEBase("application", "octet-stream")
        part.set_payload(file_stream.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f"attachment; filename={file_name}")
        message.attach(part)

        # Connect to the SMTP server and send the email
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()  # Secure the connection
            server.login(sender_email, sender_password)
            server.send_message(message)
        
        st.success("Recommendations sent to your email!")
    except Exception as e:
        st.error(f"Failed to send email: {e}")


# Streamlit UI with MongoDB integration
def main():
    # Add CSS for background image
    st.markdown(
        """
        <style>
        .stApp {
            background-color:#34cafe;
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
            height:100vh;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    st.title("Smart Yoga Asana Recommendation System")
    st.write("Get personalized yoga recommendations based on your health needs.")

    # User input for age
    age = st.number_input("Enter your age:", min_value=0, max_value=120, step=1)
    
    # User input for gender
    gender = st.selectbox("Select your gender:", ["Male", "Female", "Others"])
    
    # User input for health issue
    health_issue = st.text_area("Describe your health issue:")
    
    # User input for email address
    user_email = st.text_input("Enter your email address:")
    
    # Radio button to select file format
    file_format = st.radio("Select the file format to receive recommendations:", ["Word", "PDF"])

    # Recommendation button
    if st.button("Get Recommendations"):
        if health_issue.strip() and user_email.strip():
            matching_asanas = get_matching_asanas(age, gender, health_issue)
            if matching_asanas:
                recommendations = [f"{asana[0]}: Benefits {asana[1]}" for asana in matching_asanas]
                st.write("### Recommended Yoga Asanas:")
                for rec in recommendations:
                    st.write(f"- {rec}")
                
                # Create the file dynamically in memory
                if file_format == "Word":
                    file_stream = create_word_document(recommendations)
                    file_name = "Yoga_Recommendations.docx"
                else:
                    file_stream = create_pdf_document(recommendations)
                    file_name = "Yoga_Recommendations.pdf"

                # Send email with the file as attachment
                send_email(user_email, file_stream, file_name)
            else:
                st.write("No matching asanas found for the given input.")
        else:
            st.warning("Please provide your health issue and email address.")


# Run the Streamlit application
if __name__ == "__main__":
    main()